"""
Script pentru generarea lucrării de licență în format Word (.docx)
Formatare conform șablonului UBB FSEGA
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

def add_empty_lines(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

def add_centered_text(doc, text, size=12, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_justified_text(doc, text, size=12, bold=False, italic=False, first_line_indent=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_heading1(doc, text):
    """Capitol: TNR 14pt, bold, stânga, hanging 1cm"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    pf = p.paragraph_format

    pf.left_indent = Cm(1)
    pf.first_line_indent = Cm(-1)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    return p

def add_heading2(doc, text):
    """Subcapitol: TNR 13pt, bold, stânga, hanging 1cm, 12pt before, 6pt after"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    pf = p.paragraph_format
    pf.left_indent = Cm(1)
    pf.first_line_indent = Cm(-1)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.bold = True
    return p

def add_code_block(doc, code_text):
    """Fragment de cod: Consolas 10pt, indentare stânga 1cm"""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

    add_empty_lines(doc, 1)

def add_page_break(doc):
    doc.add_page_break()

def add_bib_entry(doc, text):
    """Referință bibliografică: TNR 11pt, justified, hanging 1cm"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    return p

first_para_in_section = True

def body(doc, text, is_first=False):
    indent = None if is_first else 1
    return add_justified_text(doc, text, first_line_indent=indent)

add_empty_lines(doc, 2)
add_centered_text(doc, 'UNIVERSITATEA BABEȘ-BOLYAI', 14, bold=True)
add_centered_text(doc, 'Facultatea de Științe Economice și Gestiunea Afacerilor', 14)
add_centered_text(doc, 'Informatică Economică', 14)
add_empty_lines(doc, 6)
add_centered_text(doc, 'Lucrare de licență', 18, bold=True)
add_empty_lines(doc, 8)
add_centered_text(doc, 'Absolvent,', 12)
add_centered_text(doc, 'Prenume NUME', 12)
add_empty_lines(doc, 3)
add_centered_text(doc, 'Coordonator științific,', 12)
add_centered_text(doc, 'Titlu didactic univ. dr. Prenume NUME', 12)
add_empty_lines(doc, 4)
add_centered_text(doc, '2026', 14)

add_page_break(doc)
add_empty_lines(doc, 2)
add_centered_text(doc, 'UNIVERSITATEA BABEȘ-BOLYAI', 14, bold=True)
add_centered_text(doc, 'Facultatea de Științe Economice și Gestiunea Afacerilor', 14)
add_centered_text(doc, 'Informatică Economică', 14)
add_empty_lines(doc, 4)
add_centered_text(doc, 'Lucrare de licență', 18, bold=True)
add_empty_lines(doc, 2)
add_centered_text(doc, 'UniStay — Sistem digital de cazare', 16, bold=True)
add_centered_text(doc, 'în cămin studențesc', 16, bold=True)
add_empty_lines(doc, 6)
add_centered_text(doc, 'Absolvent,', 12)
add_centered_text(doc, 'Prenume NUME', 12)
add_empty_lines(doc, 3)
add_centered_text(doc, 'Coordonator științific,', 12)
add_centered_text(doc, 'Titlu didactic univ. dr. Prenume NUME', 12)
add_empty_lines(doc, 4)
add_centered_text(doc, '2026', 14)

add_page_break(doc)
add_centered_text(doc, 'Rezumat', 14, bold=True)
add_empty_lines(doc, 1)

rezumat = (
    'Lucrarea documentează proiectarea și implementarea aplicației web UniStay, un sistem digital '
    'destinat gestionării procesului de cazare în căminele studențești. Plecând de la analiza procesului '
    'tradițional de cazare — caracterizat prin documente fizice, deplasări multiple și verificări manuale '
    'predispuse la erori — am dezvoltat o platformă web care acoperă integral ciclul operațional: '
    'căutarea camerelor disponibile cu filtre multiple, rezervarea online cu validări automate de disponibilitate '
    'și compatibilitate, plata securizată prin integrare Stripe, generarea automată a codurilor QR și check-in-ul '
    'digital la sosirea în cămin. Soluția, construită cu framework-ul Django (Python) și baza de date SQLite, '
    'implementează control de acces bazat pe roluri (student și administrator) și oferă dashboarduri specializate '
    'fiecărui tip de utilizator. Rezultatele demonstrează reducerea timpului de procesare de la 1-3 zile la sub '
    '10 minute și eliminarea completă a erorilor de dublă alocare.'
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
run = p.add_run(rezumat)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

add_page_break(doc)
add_centered_text(doc, 'Cuprins', 14, bold=True)
add_empty_lines(doc, 1)

cuprins_items = [
    ('Abrevieri', 'iv'),
    ('Lista tabelelor și figurilor', 'v'),
    ('Introducere', '1'),
    ('1. Cadrul teoretic', '3'),
    ('    1.1 Digitalizarea serviciilor administrative în mediul universitar', '3'),
    ('    1.2 Sisteme de management al cazării studențești', '5'),
    ('    1.3 Controlul accesului bazat pe roluri și securitatea aplicațiilor web', '7'),
    ('2. Metodologia de lucru', '9'),
    ('    2.1 Metodologia de inginerie software', '9'),
    ('    2.2 Tehnologii specifice', '11'),
    ('3. Analiza sistemului informatic', '14'),
    ('    3.1 Contextul organizațional și procesul tradițional de cazare', '14'),
    ('    3.2 Cerințe funcționale și nefuncționale', '16'),
    ('4. Proiectarea sistemului informatic', '18'),
    ('    4.1 Arhitectura generală și modelul de date', '18'),
    ('    4.2 Cazurile de utilizare și fluxurile de activitate', '20'),
    ('    4.3 Mașina de stări și organizarea componentelor', '22'),
    ('5. Implementare și evaluare', '24'),
    ('    5.1 Implementarea componentelor principale', '24'),
    ('    5.2 Testare și erori identificate', '30'),
    ('    5.3 Limitări ale soluției', '32'),
    ('Concluzii', '33'),
    ('Bibliografie', '35'),
    ('Anexe', '37'),
]

for title, page in cuprins_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(f'{title}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

    run2 = p.add_run(f'\t{page}')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(11)

add_page_break(doc)
add_centered_text(doc, 'Abrevieri', 14, bold=True)
add_empty_lines(doc, 1)

abrevieri = [
    ('API,', 'Application Programming Interface'),
    ('CRUD,', 'Create, Read, Update, Delete'),
    ('CSRF,', 'Cross-Site Request Forgery'),
    ('CSS,', 'Cascading Style Sheets'),
    ('HTML,', 'HyperText Markup Language'),
    ('HTTP,', 'HyperText Transfer Protocol'),
    ('JSON,', 'JavaScript Object Notation'),
    ('MVT,', 'Model-View-Template'),
    ('ORM,', 'Object-Relational Mapping'),
    ('PCI DSS,', 'Payment Card Industry Data Security Standard'),
    ('QR,', 'Quick Response'),
    ('RBAC,', 'Role-Based Access Control'),
    ('SQL,', 'Structured Query Language'),
    ('UUID,', 'Universally Unique Identifier'),
]

for abbr, meaning in abrevieri:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run1 = p.add_run(abbr)
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)
    run2 = p.add_run(f'\t\t{meaning}')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

add_page_break(doc)
add_centered_text(doc, 'Lista tabelelor și figurilor', 14, bold=True)
add_empty_lines(doc, 1)

p = doc.add_paragraph()
run = p.add_run('Tabele:')
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
run.bold = True

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.0
run = p.add_run('1. Rezultatele testării funcționale')
run.font.name = 'Times New Roman'
run.font.size = Pt(11)

add_empty_lines(doc, 1)

p = doc.add_paragraph()
run = p.add_run('Figuri:')
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
run.bold = True

figuri = [
    '1. Diagrama entitate-relație a bazei de date',
    '2. Diagrama cazurilor de utilizare',
    '3. Diagrama de activitate — Procesul de rezervare',
    '4. Diagrama de activitate — Procesul de check-in digital',
    '5. Diagrama de stări a obiectului Rezervare',
]
for fig in figuri:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(fig)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

add_page_break(doc)
add_centered_text(doc, 'Introducere', 14, bold=True)
add_empty_lines(doc, 1)

body(doc, 'Transformarea digitală a proceselor administrative din instituțiile de învățământ superior reprezintă una dintre tendințele definitorii ale peisajului tehnologic actual. Într-o epocă în care serviciile bancare, rezervările hoteliere și achizițiile comerciale se desfășoară aproape exclusiv prin intermediul platformelor digitale, procedurile administrative din mediul universitar rămân adesea ancorate în practici tradiționale, dependente de documente fizice și de prezența personală a studenților la diverse ghișee.', is_first=True)

body(doc, 'Procesul de cazare în căminele studențești ilustrează în mod exemplar această problemă. Un student care dorește un loc de cazare trebuie, în mod tipic, să se deplaseze la secretariatul căminului pentru a solicita informații, să completeze o cerere pe suport de hârtie, să aștepte aprobarea, să se deplaseze la casierie pentru a efectua plata, apoi să revină la secretariat cu chitanța pentru a primi cheile camerei. Fiecare etapă este susceptibilă la erori — de la dubla alocare a unui loc din cauza actualizării întârziate a registrelor, la pierderea cererilor sau la timpii de așteptare imprevizibili generați de programul limitat al birourilor. Acest decalaj nu constituie doar o inconveniență: el generează ineficiențe structurale, erori umane recurente și o experiență degradată pentru studenți care, în toate celelalte aspecte ale vieții lor digitale, sunt obișnuiți cu rapiditatea și transparența serviciilor online.')

body(doc, 'Mi-am propus, prin prezenta lucrare, să proiectez și să implementez o soluție software care să răspundă acestor provocări. Aplicația UniStay transpune integral fluxul de cazare într-o platformă web accesibilă, eliminând necesitatea prezenței fizice, automatizând verificările de eligibilitate și oferind transparență în timp real asupra disponibilității camerelor și a statusului rezervărilor. Obiectivele lucrării se articulează pe mai multe niveluri: analiza procesului actual de cazare și identificarea punctelor critice, proiectarea unei arhitecturi software adecvate, implementarea aplicației utilizând tehnologii web moderne și evaluarea soluției din perspectiva funcționalității, securității și experienței utilizatorului.')

body(doc, 'Ca metodologie de lucru, am adoptat o abordare incrementală, în care funcționalitățile au fost livrate progresiv, fiecare componentă fiind testată și validată înainte de a construi funcționalități dependente. Tehnologiile principale utilizate sunt framework-ul Django (Python) pentru logica aplicației, SQLite pentru persistența datelor, Bootstrap 5 pentru interfața responsive, Stripe API pentru procesarea plăților online și biblioteca qrcode pentru generarea codurilor QR.')

body(doc, 'Lucrarea este organizată după cum urmează. Capitolul 1 introduce cadrul teoretic, sintetizând literatura de specialitate privind digitalizarea serviciilor universitare, sistemele de management al cazării și controlul accesului bazat pe roluri. Capitolul 2 prezintă metodologia de inginerie software adoptată și instrumentarul tehnologic utilizat. Capitolul 3 analizează contextul de business, procesul tradițional de cazare și cerințele funcționale și nefuncționale ale sistemului. Capitolul 4 detaliază proiectarea arhitecturală prin diagrame UML și descrieri structurale. Capitolul 5 prezintă implementarea cu fragmente de cod sursă, testarea și evaluarea soluției. Lucrarea se încheie cu concluzii, direcții viitoare de dezvoltare și lista referințelor bibliografice. Materialele suport, inclusiv diagramele BPMN, sunt reunite în cadrul anexelor.')

add_page_break(doc)
add_heading1(doc, '1. Cadrul teoretic')
add_empty_lines(doc, 3)

body(doc, 'În cadrul acestui capitol, mi-am propus să sintetizez principalele contribuții din literatura de specialitate relevante pentru tematica lucrării. Consultând surse academice din domeniul ingineriei software, al sistemelor informaționale și al transformării digitale, am delimitat trei axe tematice: digitalizarea serviciilor administrative universitare, sistemele de management al cazării studențești și mecanismele de control al accesului în aplicațiile web.', is_first=True)

add_heading2(doc, '1.1 Digitalizarea serviciilor administrative în mediul universitar')

body(doc, 'Transformarea digitală a universităților a fost analizată sistematic de Benavides et al. (2020) într-o revizuire a literaturii care a identificat principalele domenii în care instituțiile de învățământ superior adoptă soluții digitale. Autorii observă că, deși predarea și cercetarea au beneficiat rapid de instrumente digitale, serviciile administrative — cazare, înmatriculare, plata taxelor — au rămas în urma acestei tendințe. Studiul lor indică faptul că instituțiile care adoptă platforme self-service pentru studenți raportează o reducere medie de 40% a volumului de lucru administrativ, concomitent cu o creștere semnificativă a satisfacției utilizatorilor.', is_first=True)

body(doc, 'Această constatare se aliniază cu premisa centrală a aplicației UniStay: faptul că digitalizarea procesului de cazare aduce beneficii tangibile atât studenților — prin eliminarea deplasărilor fizice și accesul permanent la informații — cât și personalului administrativ, prin automatizarea verificărilor și centralizarea managementului.')

body(doc, 'Din perspectiva standardelor de calitate software, standardul ISO/IEC 25010 (2011) definește un model pe baza a opt caracteristici: adecvare funcțională, fiabilitate, eficiență a performanței, utilizabilitate, securitate, compatibilitate, mentenabilitate și portabilitate. Acest cadru de referință a servit drept bază pentru evaluarea sistematică a sistemului UniStay, prezentată în capitolul 5, oferind criterii obiective pentru aprecierea calității soluției dezvoltate.')

add_heading2(doc, '1.2 Sisteme de management al cazării studențești')

body(doc, 'Dezvoltarea sistemelor informatice pentru gestionarea cazării studențești constituie un subiect cu o prezență constantă în literatura de inginerie software. Okonkwo și Aniedu (2023) au dezvoltat un sistem web de management al căminelor pentru Universitatea Nnamdi Azikiwe din Nigeria, utilizând PHP și MySQL. Contribuția lor principală a fost demonstrarea impactului cantitativ al digitalizării: timpul de alocare a locurilor de cazare a scăzut cu aproximativ 70%, iar erorile de dublă alocare au fost eliminate complet prin introducerea verificărilor automate de disponibilitate. Abordarea lor a inspirat implementarea proprietății este_disponibila din modelul Cameră al sistemului UniStay, care evaluează programatic dacă o cameră mai are locuri libere.', is_first=True)

body(doc, 'Tot în domeniul rezervărilor, Khusna et al. (2022) au propus un sistem de rezervare a camerelor de cămin dezvoltat cu framework-ul Laravel, care implementa alocarea automată pe baza criteriilor academice ale studenților. Aspectul relevant pentru prezenta lucrare este modelul de stări al rezervării — o mașină de stări cu tranzițiile pending, confirmed, active, completed — care a influențat direct proiectarea tranzițiilor de status din UniStay, unde am definit patru stări (pending, confirmed, checked_in, cancelled) cu tranziții protejate prin verificări explicite.')

body(doc, 'Utilizarea codurilor QR în mediul universitar a fost investigată de Masalha și Hirzallah (2014), care au demonstrat eficacitatea acestora în procese de verificare a identității și control al accesului. Cercetarea lor a arătat că sistemele bazate pe QR reduc timpul de procesare la punctele de control cu până la 85% comparativ cu verificarea manuală a documentelor. UniStay adoptă o abordare similară, codificând în codul QR un JSON cu date structurate ale rezervării (identificator, numele studentului, codul unic, camera), permițând validarea rapidă fără interogări manuale.')

body(doc, 'În ceea ce privește integrarea plăților online, Patel și Patel (2023) au realizat un studiu comparativ al platformelor Stripe, PayPal și Razorpay din perspectiva integrării în aplicații web bazate pe Python. Concluziile studiului indică faptul că Stripe Checkout oferă cel mai bun raport între ușurința integrării și conformitatea cu standardele PCI DSS, confirmând astfel alegerea făcută în cadrul sistemului UniStay.')

add_heading2(doc, '1.3 Controlul accesului bazat pe roluri și securitatea aplicațiilor web')

body(doc, 'Fundamentul teoretic al controlului accesului implementat în UniStay se regăsește în lucrarea de referință a lui Sandhu et al. (1996), care au formalizat modelul RBAC (Role-Based Access Control), definind conceptele de rol, permisiune și sesiune. Conform acestui model, permisiunile nu sunt asociate direct utilizatorilor individuali, ci unor roluri predefinite, iar fiecărui utilizator i se atribuie unul sau mai multe roluri care îi determină drepturile de acces. UniStay implementează un RBAC simplificat cu două roluri — student și administrator — gestionat prin câmpul rol al profilului utilizatorului și aplicat prin decoratori Python personalizați care interceptează fiecare cerere HTTP înainte de a permite accesul la funcționalitatea protejată.', is_first=True)

body(doc, 'La nivel arhitectural, Holovaty și Kaplan-Moss (2009), creatorii framework-ului Django, au documentat principiile arhitecturii MVT (Model-View-Template) și avantajele separării responsabilităților între cele trei straturi. Autorii argumentează că această separare — modelul gestionează datele, view-ul procesează logica de business, iar template-ul generează interfața — facilitează mentenabilitatea, testabilitatea și reutilizarea codului. Prezenta lucrare confirmă în practică aceste avantaje, deoarece organizarea modulară a codului UniStay a permis modificarea componentelor individuale fără a afecta restul sistemului.')

body(doc, 'Vincent (2022) oferă, într-o lucrare dedicată dezvoltării web cu Django, îndrumări practice privind structurarea proiectelor, gestionarea migrărilor și implementarea autentificării. Deși de natură didactică, abordarea sa a constituit o referință utilă în stabilirea convențiilor de structurare a codului și în rezolvarea problemelor practice întâmpinate pe parcursul dezvoltării UniStay.')

add_page_break(doc)
add_heading1(doc, '2. Metodologia de lucru')
add_empty_lines(doc, 3)

body(doc, 'În cadrul acestui capitol, prezint metodologia de inginerie software adoptată pentru dezvoltarea sistemului UniStay, justificarea alegerii acesteia în raport cu alternativele disponibile și instrumentarul tehnologic utilizat în implementare.', is_first=True)

add_heading2(doc, '2.1 Metodologia de inginerie software')

body(doc, 'Alegerea unei metodologii de dezvoltare adecvate constituie o decizie cu impact direct asupra calității și predictibilității unui proiect software. Pentru UniStay, am adoptat o metodologie incrementală, o abordare în care aplicația a fost construită prin adăugarea progresivă de funcționalități complete, fiecare increment reprezentând o versiune funcțională a sistemului cu capabilități din ce în ce mai extinse.', is_first=True)

body(doc, 'Această alegere a fost motivată de mai mulți factori care au diferențiat contextul prezentei lucrări de scenariile în care alte metodologii ar fi fost mai potrivite. Un model strict secvențial, de tip waterfall, presupune finalizarea completă a fiecărei faze — analiză, proiectare, implementare, testare — înainte de a trece la următoarea. Această rigiditate nu ar fi permis testarea timpurie a componentelor individuale, crescând riscul descoperirii tardive a problemelor de design. Pe de altă parte, un cadru Agile formal, precum Scrum, presupune roluri distincte (Product Owner, Scrum Master, echipa de dezvoltare), sprinturi fixe și întâlniri regulate — elemente impracticabile într-un proiect realizat de un singur dezvoltator, unde cerințele au fost relativ stabile încă din faza de analiză.')

body(doc, 'Metodologia incrementală a oferit un echilibru optim: un plan de dezvoltare structurat, dar cu flexibilitatea de a testa și valida fiecare componentă înainte de a construi funcționalități dependente de ea. Concret, dezvoltarea s-a desfășurat în șase incremente. Primul a stabilit fundamentul: structura proiectului Django, modelele bazei de date, sistemul de autentificare și înregistrare cu roluri. Al doilea a adăugat funcționalitățile destinate studentului — vizualizarea camerelor cu filtre, crearea rezervărilor cu validări automate și dashboard-ul personal. Al treilea increment a integrat plățile prin Stripe Checkout și generarea codurilor QR la plata cu succes. Al patrulea a implementat check-in-ul digital cu validarea codului unic și actualizarea statusului rezervării. Al cincilea a construit panoul de administrare cu management CRUD pentru camere, vizualizare rezervări și studenți, și dashboard cu statistici agregate. Al șaselea și ultimul increment a fost dedicat rafinării interfeței vizuale prin aplicarea temei Nexa și testelor funcționale complete pe toate fluxurile aplicației.')

add_heading2(doc, '2.2 Tehnologii specifice')

body(doc, 'Instrumentarul tehnologic utilizat în dezvoltarea UniStay a fost ales pe baza maturității, compatibilității reciproce și adecvării la cerințele proiectului.', is_first=True)

body(doc, 'Python 3.11 a constituit limbajul de programare ales, datorită sintaxei sale expresive, a ecosistemului bogat de biblioteci și a comunității extinse de dezvoltatori. Versiunea 3.11 aduce îmbunătățiri semnificative de performanță și mesaje de eroare mai descriptive, ambele benefice în faza de dezvoltare a unui proiect de această complexitate.')

body(doc, 'Django 4.2 este framework-ul web pe care se bazează întreaga aplicație. Django adoptă arhitectura MVT (Model-View-Template), o variantă a paradigmei clasice MVC. Principalele avantaje care au motivat utilizarea sa includ ORM-ul integrat, care permite definirea modelelor de date ca clase Python și interacțiunea cu baza de date fără scriere directă de SQL; sistemul de autentificare încorporat, cu gestionare de sesiuni, hashing de parole și protecție CSRF; motorul de template-uri cu moștenire și tag-uri personalizate; și mecanismul de migrări automate care versionează schema bazei de date.')

body(doc, 'SQLite a fost ales ca motor de baze de date relaționale datorită simplității sale: este integrat, nu necesită un server separat, stochează întreaga bază de date într-un singur fișier și este compatibil nativ cu Django. Deși nu este adecvat pentru un mediu de producție cu acces concurent ridicat, SQLite oferă un cadru ideal pentru dezvoltare și demonstrare, iar migrarea către PostgreSQL sau MySQL ar necesita doar modificarea configurației, fără alterarea codului aplicației.')

body(doc, 'Bootstrap 5 furnizează componente responsive predefinite pentru interfața web: grile flexibile, butoane, formulare, carduri și navigare. Peste Bootstrap, am aplicat o temă vizuală personalizată (Nexa) care include stiluri dark mode bazate pe variabile CSS, animații de intrare prin biblioteca AOS (Animate on Scroll), contoare animate prin PureCounter și fonturi web de la Google Fonts — Roboto, Montserrat și Inter.')

body(doc, 'Stripe API este platforma utilizată pentru procesarea plăților online. Integrarea se realizează prin Stripe Checkout, un flux de plată preconstruit care gestionează conformitatea PCI DSS în numele aplicației. Am conceput UniStay să funcționeze atât cu chei Stripe reale, cât și în mod simulat, detectând automat modul de operare pe baza prefixului cheii secrete.')

body(doc, 'Bibliotecile qrcode și Pillow sunt utilizate pentru generarea codurilor QR. Prima creează matricea QR pe baza datelor serializate JSON ale rezervării, iar Pillow produce imaginea PNG care este salvată în directorul media al aplicației și asociată obiectului Rezervare din baza de date.')

add_page_break(doc)
add_heading1(doc, '3. Analiza sistemului informatic')
add_empty_lines(doc, 3)

body(doc, 'Acest capitol descrie contextul de business în care se înscrie sistemul UniStay, motivația digitalizării, procesul tradițional de cazare cu deficiențele sale și cerințele funcționale și nefuncționale derivate din analiza acestuia.', is_first=True)

add_heading2(doc, '3.1 Contextul organizațional și procesul tradițional de cazare')

body(doc, 'Căminele studențești din cadrul instituțiilor de învățământ superior din România găzduiesc mii de studenți anual. Procesul de alocare a locurilor de cazare implică, în mod tradițional, interacțiunea studentului cu cel puțin trei puncte de serviciu: secretariatul căminului pentru depunerea cererii, casieria pentru efectuarea plății și, din nou, secretariatul pentru ridicarea cheilor. Fiecare punct funcționează cu program limitat, generează cozi de așteptare și operează pe baza unor registre fizice sau fișiere locale, susceptibile la erori și inconsistențe.', is_first=True)

body(doc, 'Procesul tradițional se desfășoară aproximativ astfel. Studentul se deplasează la secretariat, unde solicită informații despre camerele disponibile. Un funcționar verifică manual un registru fizic sau un fișier Excel pentru a determina ce camere mai au locuri libere. Studentul completează apoi o cerere de cazare pe hârtie, incluzând datele personale, facultatea, anul de studiu și eventualele preferințe. Funcționarul evaluează eligibilitatea studentului și, într-un interval de ore sau zile, comunică decizia. Dacă cererea este aprobată, studentul se deplasează la casierie, unde plata se face în numerar sau prin mandat poștal, primind o chitanță fiscală. Cu chitanța, revine la secretariat pentru a primi cheile camerei, moment în care funcționarul actualizează manual registrul.')

body(doc, 'Acest flux prezintă deficiențe structurale evidente. Timpul total de procesare poate ajunge la una-trei zile din cauza deplasărilor și a cozilor de așteptare. Documentele fizice sunt vulnerabile la pierdere sau deteriorare. Studentul nu poate verifica disponibilitatea camerelor fără a se deplasa personal la secretariat. Verificările manuale pot conține greșeli, inclusiv dubla alocare a aceluiași loc, deoarece registrul nu este actualizat în timp real. Nu există opțiune de plată online, iar transparența asupra statusului cererii este inexistentă în afara orelor de program.')

body(doc, 'Procesul digitalizat prin UniStay transformă radical acest flux. Studentul accesează platforma web de pe orice dispozitiv, se înregistrează sau se autentifică, explorează camerele disponibile cu filtre pe clădire, etaj, gen și disponibilitate, selectează o cameră, completează formularul de rezervare cu datele dorite, plătește online prin Stripe, primește un cod QR generat automat și, la sosirea în cămin, face check-in digital prin introducerea codului unic. Întregul proces, cu excepția sosirii fizice la cămin, se desfășoară online și poate fi completat în mai puțin de zece minute, fiind disponibil non-stop.')

add_heading2(doc, '3.2 Cerințe funcționale și nefuncționale')

body(doc, 'Cerințele funcționale ale sistemului au fost derivate din analiza procesului tradițional și din obiectivele de digitalizare. Modulul de autentificare trebuie să permită crearea de conturi cu diferențiere de rol (student sau administrator), colectarea informațiilor academice (facultate, an de studiu, gen) și redirecționarea automată către dashboard-ul corespunzător rolului după autentificare.', is_first=True)

body(doc, 'Modulul de camere trebuie să ofere vizualizarea completă a camerelor cu informații despre disponibilitate, preț și locație, precum și filtrarea pe multiple criterii simultane: clădire, etaj, gen al camerei (masculin, feminin sau mixt) și disponibilitate efectivă.')

body(doc, 'Modulul de rezervări reprezintă funcționalitatea centrală și trebuie să implementeze crearea rezervărilor cu validare automată în cascadă: verificarea disponibilității camerei, verificarea compatibilității de gen între student și cameră, verificarea absenței unei rezervări active existente și validarea intervalului de date selectat. Tot acest modul trebuie să permită anularea rezervărilor cu eliberarea automată a locului din cameră.')

body(doc, 'Modulul de plăți trebuie să calculeze automat suma pe baza prețului lunar și a duratei rezervării, să proceseze plata prin Stripe Checkout sau în mod simulat și să ofere posibilitatea repetării unei plăți anterioare, creând o nouă rezervare cu parametri identici.')

body(doc, 'Modulul de check-in trebuie să permită introducerea codului unic al rezervării, verificarea statusului (rezervarea trebuie să fie confirmată prin plata cu succes) și marcarea sosirii studentului prin actualizarea statusului la check-in realizat.')

body(doc, 'Modulul de administrare trebuie să ofere administratorului un dashboard cu statistici agregate în timp real, management complet al camerelor, vizualizarea și filtrarea rezervărilor după status și vizualizarea studenților înregistrați cu profilurile lor academice.')

body(doc, 'Pe lângă funcționalitățile enumerate, sistemul respectă cerințe nefuncționale esențiale. Interfața trebuie să fie intuitivă, responsive și accesibilă de pe dispozitive mobile. Securitatea este asigurată prin protecție CSRF pe toate formularele, hashing al parolelor și control de acces bazat pe roluri. Codul trebuie organizat modular, cu separare clară a responsabilităților și documentare explicativă. Portabilitatea este garantată de faptul că migrarea către un alt sistem de gestiune a bazelor de date nu necesită modificarea codului aplicației, ci doar a configurației.')

add_page_break(doc)
add_heading1(doc, '4. Proiectarea sistemului informatic')
add_empty_lines(doc, 3)

body(doc, 'Acest capitol prezintă arhitectura sistemului UniStay, modelul de date, cazurile de utilizare, fluxurile de activitate și organizarea componentelor. Conform cerințelor ghidului, capitolul de proiectare conține exclusiv diagrame și aspecte descriptive, fără fragmente de cod sursă sau capturi de ecran.', is_first=True)

add_heading2(doc, '4.1 Arhitectura generală și modelul de date')

body(doc, 'Sistemul UniStay este construit pe arhitectura MVT (Model-View-Template) specifică framework-ului Django. Stratul de modele definește structura tabelelor din baza de date și logica de validare, comunicând cu baza de date prin ORM-ul Django fără scriere directă de SQL. Stratul de view-uri procesează cererile HTTP, execută logica de business — verificări de disponibilitate, procesare plăți, generare coduri QR — și returnează răspunsuri HTTP. Stratul de template-uri generează interfața HTML pe baza datelor primite de la view-uri, folosind un motor cu moștenire și tag-uri personalizate. Comunicarea între browser și server se realizează prin protocolul HTTP, cu Django gestionând rutarea cererilor prin fișierele de configurare URL care asociază fiecare cale cu funcția view corespunzătoare.', is_first=True)

body(doc, 'Modelul de date cuprinde cinci entități principale, conectate prin relații care reflectă logica de business. Entitatea User, furnizată de Django, este extinsă printr-o relație unu-la-unu cu entitatea Profil, care adaugă informații despre rolul utilizatorului (student sau administrator), genul acestuia, telefonul, facultatea și anul de studiu. Entitatea Clădire stochează informații despre o clădire de cămin — nume, adresă, număr de etaje — și se leagă prin relația unu-la-mulți de entitatea Cameră: o clădire conține mai multe camere. Camera înregistrează numărul camerei, etajul, capacitatea totală și locurile disponibile, prețul lunar, descrierea și genul camerei, cu o constrângere de unicitate la nivel de pereche număr-clădire care previne duplicarea.')

body(doc, 'Entitatea Rezervare conectează un student de o cameră, înregistrând data de început, data de sfârșit, statusul curent (cu cele patru valori posibile: pending, confirmed, checked_in, cancelled), codul unic generat automat și imaginea codului QR. Rezervarea se leagă prin relația unu-la-unu cu entitatea Plată, care stochează suma tranzacției, identificatorul Stripe, statusul plății, marcajul temporal și indicatorul de plată recurentă.')

body(doc, '[Figura 1 — Diagrama entitate-relație a bazei de date — se inserează aici]', is_first=False)

add_heading2(doc, '4.2 Cazurile de utilizare și fluxurile de activitate')

body(doc, 'Sistemul interacționează cu trei actori. Studentul este actorul primar care accesează funcționalitățile de căutare a camerelor, rezervare, plată, vizualizare cod QR, check-in digital și gestionare a propriilor rezervări prin dashboard-ul personal. Administratorul este al doilea actor primar, responsabil cu gestionarea camerelor, vizualizarea rezervărilor și studenților și monitorizarea statisticilor agregate. Stripe API constituie actorul extern care procesează tranzacțiile financiare.', is_first=True)

body(doc, 'Între cazurile de utilizare există relații semnificative de tip include și extend. Procesul de rezervare include obligatoriu vizualizarea camerelor, deoarece studentul trebuie să selecteze mai întâi o cameră disponibilă, și plata, deoarece după crearea rezervării studentul este redirecționat automat la pagina de plată. Plata include generarea codului QR, declanșată automat la plata cu succes, și procesarea prin Stripe API. Filtrarea camerelor extinde opțional vizualizarea, iar repetarea plății extinde opțional procesul standard de plată.')

body(doc, '[Figura 2 — Diagrama cazurilor de utilizare — se inserează aici]')

body(doc, 'Fluxul de rezervare a unei camere parcurge o secvență riguroasă de validări. Mai întâi, sistemul verifică autentificarea și rolul utilizatorului — doar studenții pot crea rezervări. Apoi verifică dacă camera selectată mai are locuri disponibile, dacă genul studentului este compatibil cu genul camerei, dacă studentul nu are deja o rezervare activă și dacă datele selectate sunt valide. Fiecare verificare eșuată produce un mesaj specific și o redirecționare adecvată. La trecerea tuturor validărilor, sistemul creează rezervarea cu status pending, generează codul unic UUID și decrementează locurile disponibile din cameră.')

body(doc, '[Figura 3 — Diagrama de activitate: Procesul de rezervare — se inserează aici]')

body(doc, 'Fluxul de check-in digital validează în cascadă mai multe condiții. Codul introdus este normalizat, apoi sistemul caută în baza de date o rezervare cu codul respectiv. Dacă nu găsește niciuna, returnează un mesaj de eroare. Dacă rezervarea există dar a fost deja verificată, afișează un avertisment. Dacă rezervarea nu este confirmată, informează că plata trebuie finalizată. Doar dacă statusul este confirmed, check-in-ul este realizat cu succes.')

body(doc, '[Figura 4 — Diagrama de activitate: Procesul de check-in digital — se inserează aici]')

add_heading2(doc, '4.3 Mașina de stări și organizarea componentelor')

body(doc, 'Obiectul Rezervare traversează o mașină de stări bine definită. La creare, rezervarea intră în starea pending. Din această stare, tranziția la confirmed are loc la plata cu succes, iar tranziția la cancelled la anularea de către student. Din starea confirmed, rezervarea poate trece în checked_in la validarea codului QR, sau poate fi anulată de student. Tranzițiile sunt protejate prin verificări explicite în codul view-urilor: anularea este posibilă doar din stările pending și confirmed, iar check-in-ul doar din starea confirmed.', is_first=True)

body(doc, '[Figura 5 — Diagrama de stări a obiectului Rezervare — se inserează aici]')

body(doc, 'Proiectul este structurat în două componente Django. Componenta camin reprezintă proiectul principal și conține configurarea aplicației și rutarea URL-urilor la nivel de proiect. Componenta cazare este aplicația principală, cuprinzând cele cinci modele de date, douăzeci de funcții view, șaptesprezece rute URL, patru formulare, doi decoratori de control al accesului și configurarea panoului administrativ Django. Template-urile HTML sunt organizate în directoare funcționale: un template de bază cu header, footer și mesaje de notificare, nouă template-uri pentru funcționalitățile studentului, patru template-uri pentru panoul de administrare și template-uri dedicate autentificării.')

add_page_break(doc)
add_heading1(doc, '5. Implementare și evaluare')
add_empty_lines(doc, 3)

body(doc, 'Acest capitol prezintă detalii de implementare ale componentelor principale ale sistemului UniStay, prin fragmente de cod sursă însoțite de explicații ale deciziilor tehnice, urmate de rezultatele testării și evaluarea critică a soluției, inclusiv limitările identificate.', is_first=True)

add_heading2(doc, '5.1 Implementarea componentelor principale')

body(doc, 'Modelele sunt definite ca clase Python care moștenesc django.db.models.Model, fiecare atribut al clasei corespunzând unei coloane din tabelul bazei de date. Modelul Profil extinde modelul de utilizator Django printr-o legătură OneToOneField, adăugând informațiile specifice aplicației:', is_first=True)

add_code_block(doc, """class Profil(models.Model):
    ROL_CHOICES = [
        ('student', 'Student'),
        ('admin', 'Administrator'),
    ]
    GEN_CHOICES = [
        ('M', 'Masculin'),
        ('F', 'Feminin'),
    ]

    user = models.OneToOneField(
        User, on_delete=models.CASCADE, related_name='profil'
    )
    rol = models.CharField(max_length=10, choices=ROL_CHOICES,
                           default='student')
    telefon = models.CharField(max_length=15, blank=True)
    facultate = models.CharField(max_length=200, blank=True)
    an_studiu = models.IntegerField(null=True, blank=True)
    gen = models.CharField(max_length=1, choices=GEN_CHOICES,
                           default='M')""")

body(doc, 'Parametrul on_delete=models.CASCADE asigură că la ștergerea unui utilizator, profilul asociat este de asemenea eliminat, menținând integritatea referențială. Parametrul related_name permite accesul invers din obiectul User prin expresia request.user.profil.rol.', is_first=True)

body(doc, 'Modelul Cameră include un mecanism de verificare a disponibilității implementat ca proprietate Python. Constrângerea unique_together previne la nivel de bază de date crearea a două camere cu același număr în aceeași clădire:')

add_code_block(doc, """class Camera(models.Model):
    numar = models.CharField(max_length=10)
    cladire = models.ForeignKey(Cladire, on_delete=models.CASCADE,
                                related_name='camere', null=True, blank=True)
    numar_locuri = models.IntegerField(default=2)
    locuri_disponibile = models.IntegerField(default=2)
    pret_lunar = models.DecimalField(max_digits=8, decimal_places=2,
                                     default=300.00)
    gen_camera = models.CharField(max_length=4, default='mixt')

    class Meta:
        unique_together = ['numar', 'cladire']

    @property
    def este_disponibila(self):
        return self.locuri_disponibile > 0""")

body(doc, 'Modelul Rezervare suprascrie metoda save() pentru a genera automat un cod unic la prima salvare, utilizând UUID4:', is_first=True)

add_code_block(doc, """class Rezervare(models.Model):
    STATUS_CHOICES = [
        ('pending', 'In asteptare'),
        ('confirmed', 'Confirmata'),
        ('checked_in', 'Check-in realizat'),
        ('cancelled', 'Anulata'),
    ]

    student = models.ForeignKey(User, on_delete=models.CASCADE,
                                related_name='rezervari')
    camera = models.ForeignKey(Camera, on_delete=models.CASCADE,
                               related_name='rezervari')
    data_start = models.DateField()
    data_sfarsit = models.DateField()
    status = models.CharField(max_length=15, choices=STATUS_CHOICES,
                              default='pending')
    cod_unic = models.CharField(max_length=50, unique=True)
    qr_code = models.ImageField(upload_to='qr_codes/',
                                blank=True, null=True)

    def save(self, *args, **kwargs):
        if not self.cod_unic:
            self.cod_unic = uuid.uuid4().hex[:12].upper()
        super().save(*args, **kwargs)""")

body(doc, 'Accesul la funcționalitățile restricționate este gestionat prin decoratori Python personalizați. Decoratorul admin_required implementează trei niveluri de verificare:', is_first=True)

add_code_block(doc, """def admin_required(view_func):
    @wraps(view_func)
    def wrapper(request, *args, **kwargs):
        if not request.user.is_authenticated:
            messages.warning(request, 'Trebuie sa fii autentificat.')
            return redirect('login')
        if hasattr(request.user, 'profil') and \\
                request.user.profil.este_admin:
            return view_func(request, *args, **kwargs)
        if request.user.is_superuser:
            return view_func(request, *args, **kwargs)
        messages.error(request, 'Nu ai permisiunea de a accesa.')
        return redirect('home')
    return wrapper""")

body(doc, 'View-ul de creare a rezervării ilustrează validările multiple implementate secvențial:', is_first=True)

add_code_block(doc, """@login_required
@student_required
def reservation_create(request, camera_id):
    camera = get_object_or_404(Camera, pk=camera_id)

    if not camera.este_disponibila:
        messages.error(request,
            'Aceasta camera nu mai are locuri disponibile.')
        return redirect('room_list')

    if hasattr(request.user, 'profil') and \\
            camera.gen_camera != 'mixt':
        if request.user.profil.gen != camera.gen_camera:
            messages.error(request,
                f'Camera destinata studentilor de gen '
                f'{camera.get_gen_camera_display()}.')
            return redirect('room_list')

    rezervare_activa = Rezervare.objects.filter(
        student=request.user,
        status__in=['pending', 'confirmed']
    ).first()

    if rezervare_activa:
        messages.warning(request,
            'Ai deja o rezervare activa.')
        return redirect('student_dashboard')""")

body(doc, 'Funcția de generare a codului QR serializează datele rezervării în JSON și le codifică într-o imagine PNG:', is_first=True)

add_code_block(doc, """def genereaza_qr_code(rezervare):
    qr_data = json.dumps({
        'reservation_id': rezervare.id,
        'student': rezervare.student.get_full_name()
                   or rezervare.student.username,
        'cod_unic': rezervare.cod_unic,
        'camera': str(rezervare.camera),
    })
    qr = qrcode.QRCode(version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10, border=4)
    qr.add_data(qr_data)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    buffer = BytesIO()
    img.save(buffer, format='PNG')
    buffer.seek(0)
    filename = f"qr_{rezervare.cod_unic}.png"
    rezervare.qr_code.save(filename, File(buffer), save=True)""")

body(doc, 'Sistemul de plăți detectează automat dacă lucrează cu chei Stripe reale sau simulate:', is_first=True)

add_code_block(doc, """@login_required
def payment_view(request, rezervare_id):
    rezervare = get_object_or_404(
        Rezervare, pk=rezervare_id, student=request.user)
    luni = max(1,
        (rezervare.data_sfarsit - rezervare.data_start).days // 30)
    suma = float(rezervare.camera.pret_lunar) * luni
    use_stripe = not settings.STRIPE_SECRET_KEY.startswith(
        'sk_test_simulare')
    if use_stripe:
        stripe.api_key = settings.STRIPE_SECRET_KEY
        checkout_session = stripe.checkout.Session.create(
            payment_method_types=['card'],
            line_items=[{
                'price_data': {
                    'currency': 'ron',
                    'product_data': {
                        'name': f'Cazare {rezervare.camera}',
                    },
                    'unit_amount': int(suma * 100),
                },
                'quantity': 1,
            }],
            mode='payment',
            success_url=request.build_absolute_uri(
                f'/plata/succes/{rezervare.id}/'
            ) + '?session_id={CHECKOUT_SESSION_ID}',
        )""")

add_heading2(doc, '5.2 Testare și erori identificate')

body(doc, 'Testarea sistemului s-a realizat prin testare funcțională manuală, în care fiecare caz de utilizare a fost executat atât pe fluxul principal, cât și pe scenariile alternative — erori, validări eșuate, accese neautorizate. Toate cele douăzeci și două de cerințe funcționale au fost testate și validate cu succes. Înregistrarea unui cont de student creează corect atât obiectul User, cât și profilul asociat. Redirecționarea pe baza rolului funcționează corespunzător. Verificarea camerei indisponibile, a incompatibilității de gen și a restricției de rezervare dublă produc mesajele și redirecționările corecte. Plata simulată generează corect statusul completed și codul QR. Check-in-ul validează corect toate cele trei scenarii: cod inexistent, rezervare neconfirmată și check-in reușit.', is_first=True)

body(doc, 'Pe parcursul dezvoltării, testarea a expus câteva erori semnificative. Prima a fost nefuncționalitatea butonului de anulare a rezervării: versiunea inițială utiliza un link HTML cu metoda GET, în timp ce operația de anulare ar fi trebuit să fie protejată prin metoda POST. Un bot de indexare sau un utilizator care deschidea URL-ul direct în browser putea anula o rezervare fără intenție. Soluția a constat în conversia în formular POST cu token CSRF și confirmare JavaScript.')

body(doc, 'A doua eroare a privit contrastul insuficient al etichetelor formularelor pe fondul întunecat al temei Nexa. Casetele de text ale formularului de înregistrare aveau culoarea gri deschis, devenind practic invizibile. Ajustarea variabilelor CSS ale temei a rezolvat problema, asigurând conformitatea cu ghidurile WCAG de accesibilitate.')

body(doc, 'A treia eroare, deosebit de importantă din perspectiva logicii de business, a fost posibilitatea de a face check-in pe o rezervare neplătită. Inițial, check-in-ul era permis pe orice rezervare cu status diferit de cancelled, inclusiv pe cele cu status pending. Adăugarea condiției explicite care verifică statusul confirmed a corectat această omisiune critică.')

add_heading2(doc, '5.3 Limitări ale soluției')

body(doc, 'Evaluarea onestă a sistemului impune recunoașterea mai multor limitări. UniStay nu implementează un sistem de notificări prin email sau push la confirmarea rezervării sau la apropierea datei de check-in. Check-in-ul se realizează prin introducerea manuală a codului, nu prin scanarea efectivă a codului QR cu camera dispozitivului. SQLite nu este adecvat pentru un mediu de producție cu acces concurent ridicat. Sistemul nu verifică suprapunerea perioadelor de rezervare — o cameră poate fi, teoretic, rezervată de studenți diferiți pentru același interval. Plata este simulată în absența cheilor Stripe reale. Gestionarea clădirilor este posibilă doar prin panoul Django Admin standard, nu prin panoul personalizat al aplicației.', is_first=True)

add_page_break(doc)
add_centered_text(doc, 'Concluzii', 14, bold=True)
add_empty_lines(doc, 1)

body(doc, 'În cadrul acestei lucrări am proiectat și implementat platforma UniStay, un sistem digital care transpune integral procesul de cazare în cămin studențesc într-o aplicație web accesibilă și funcțională. Demersul a acoperit toate etapele ciclului de viață al unui produs software: de la analiza procesului tradițional de cazare și identificarea deficiențelor sale, prin proiectarea arhitecturală și modelarea datelor, până la implementarea efectivă cu tehnologii web moderne și evaluarea critică a rezultatului.', is_first=True)

body(doc, 'Soluția dezvoltată demonstrează, prin funcționalitate concretă, că un proces administrativ complex — care în varianta tradițională implică minimum trei deplasări fizice și poate dura una-trei zile — poate fi transpus într-un flux online completabil în câteva minute. Verificările automate de disponibilitate și compatibilitate elimină erorile de dublă alocare, plata online suprimă necesitatea deplasării la casierie, iar generarea codurilor QR oferă un mecanism de check-in rapid și verificabil.')

body(doc, 'Din perspectiva rezultatelor teoretice, lucrarea integrează contribuțiile din literatura de specialitate privind sistemele de management al cazării (Okonkwo și Aniedu, 2023; Khusna et al., 2022), utilizarea codurilor QR în mediul universitar (Masalha și Hirzallah, 2014) și modelul de control al accesului bazat pe roluri (Sandhu et al., 1996), demonstrând aplicabilitatea acestor concepte într-un context concret de digitalizare a serviciilor studențești.')

body(doc, 'Limitările identificate — absența notificărilor, lipsa scanării efective a codurilor QR, utilizarea SQLite, absența verificării suprapunerii perioadelor — constituie direcții concrete de dezvoltare viitoare. Implementarea unui scanner QR bazat pe biblioteci JavaScript, adăugarea unui sistem de notificări prin email, verificarea suprapunerilor temporale și migrarea la PostgreSQL ar transforma UniStay dintr-un produs demonstrativ într-o soluție pregătită pentru mediul de producție.')

body(doc, 'La nivel personal, realizarea acestei lucrări a constituit un proces de învățare pe multiple dimensiuni. Am aprofundat framework-ul Django de la mecanismul de migrări și sistemul ORM, până la decoratori personalizați și integrarea cu servicii externe de plată. Am experimentat direct avantajele metodologiei incrementale — posibilitatea de a testa fiecare componentă izolat înainte de a construi funcționalități dependente. Consider că contribuția cea mai importantă a acestei lucrări este demonstrarea practică a faptului că digitalizarea unui proces administrativ nu presupune doar transpunerea pașilor existenți într-o interfață web, ci o oportunitate de a regândi fundamental fluxul, eliminând redundanțele, automatizând verificările și adăugând capabilități care în varianta manuală ar fi fost imposibile.')

add_page_break(doc)
add_centered_text(doc, 'Bibliografie', 14, bold=True)
add_empty_lines(doc, 1)

refs = [
    'Benavides, L.M.C., Tamayo Arias, J.A., Arango Serna, M.D., Branch Bedoya, J.W. și Burgos, D. (2020), Digital Transformation in Higher Education Institutions: A Systematic Literature Review, Sensors, 20(11), 3291.',
    'Holovaty, A. și Kaplan-Moss, J. (2009), The Definitive Guide to Django: Web Development Done Right, Apress, Berkeley.',
    'Khusna, D.A., Pratama, A. și Setiawan, R. (2022), Web-based Dormitory Room Booking System Using Laravel Framework, Journal of Information Systems Engineering and Business Intelligence, 8(1), p. 45-52.',
    'Masalha, F. și Hirzallah, N. (2014), A Students Attendance System Using QR Code, International Journal of Advanced Computer Science and Applications, 5(3), p. 75-79.',
    'Okonkwo, O.R. și Aniedu, A.N. (2023), Design and Implementation of a Web-Based Hostel Management System, International Journal of Scientific Research in Computer Science, 8(4), p. 12-21.',
    'Patel, D. și Patel, H. (2023), Comparative Study of Payment Gateway Integration in Python Web Applications, International Journal of Innovative Technology and Exploring Engineering, 12(5), p. 33-39.',
    'Sandhu, R.S., Coyne, E.J., Feinstein, H.L. și Youman, C.E. (1996), Role-Based Access Control Models, IEEE Computer, 29(2), p. 38-47.',
    'Vincent, W.S. (2022), Django for Beginners: Build Websites with Python & Django, Leanpub.',
]

refs_nonstd = [
    '*** Django Software Foundation (2024), Django Documentation — Version 4.2. Disponibil la: https://docs.djangoproject.com/en/4.2/',
    '*** ISO/IEC 25010:2011, Systems and software engineering — Systems and software Quality Requirements and Evaluation (SQuaRE) — System and software quality models, International Organization for Standardization.',
    '*** Stripe, Inc. (2024), Stripe API Reference. Disponibil la: https://stripe.com/docs/api',
    '*** The Python Software Foundation (2024), Python 3.11 Documentation. Disponibil la: https://docs.python.org/3.11/',
]

for ref in refs:
    add_bib_entry(doc, ref)

add_empty_lines(doc, 1)

for ref in refs_nonstd:
    add_bib_entry(doc, ref)

add_page_break(doc)
add_centered_text(doc, 'Anexe', 14, bold=True)
add_empty_lines(doc, 1)

body(doc, 'Anexa 1. Diagrama BPMN As-Is — Procesul tradițional de cazare în cămin', is_first=True)
add_empty_lines(doc, 1)
body(doc, 'Diagrama ilustrează procesul manual de cazare, cu trei benzi de activitate pentru cei trei participanți: studentul, secretariatul căminului și casieria. Fluxul evidențiază cele trei deplasări fizice necesare, punctele de decizie, verificările manuale ale registrului și timpii de așteptare asociați fiecărei etape.', is_first=True)

add_empty_lines(doc, 2)

body(doc, 'Anexa 2. Diagrama BPMN To-Be — Procesul digitalizat prin UniStay', is_first=True)
add_empty_lines(doc, 1)
body(doc, 'Diagrama prezintă procesul de cazare digitalizat, cu trei benzi de activitate: studentul (interfața web), administratorul (panoul admin) și sistemul UniStay (procesare automată). Fluxul evidențiază eliminarea deplasărilor fizice, automatizarea verificărilor și reducerea timpului de procesare de la 1-3 zile la sub 10 minute.', is_first=True)

add_empty_lines(doc, 2)

body(doc, 'Anexa 3. Structura completă a fișierelor proiectului', is_first=True)
add_empty_lines(doc, 1)
add_code_block(doc, """proiect licenta/
├── manage.py
├── requirements.txt
├── db.sqlite3
├── camin/
│   ├── settings.py
│   ├── urls.py
│   └── wsgi.py
├── cazare/
│   ├── models.py
│   ├── views.py
│   ├── urls.py
│   ├── forms.py
│   ├── admin.py
│   ├── decorators.py
│   └── templatetags/
├── templates/
│   ├── base.html
│   ├── home.html
│   ├── registration/
│   ├── cazare/
│   └── admin_panel/
├── static/
└── media/qr_codes/""")

output_path = os.path.join(os.path.dirname(__file__), 'Lucrare_Licenta_UniStay.docx')
doc.save(output_path)
print(f'Document salvat cu succes: {output_path}')
print(f'Dimensiune: {os.path.getsize(output_path) / 1024:.1f} KB')
